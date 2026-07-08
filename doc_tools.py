import os
from typing import Tuple, List
import numpy as np

from langchain_core.tools import tool
from docx import Document
from docx.shared import Pt

from vec_search import vector_store

GLOBAL_DOC = Document()
DOC_PATH = os.path.join("generated","demo.docx")

@tool
def write_markdown(content: str) -> None:
    """
    Writes LLM generated todo list content in a md file
    :param content:
    :return:
    """
    with open(os.path.join("task","TODO.md"), "w") as md:
        md.write(content)

@tool
def write_paragraph(para: str) -> str:
    """
    writes paragraph in the doc
    :param para: paragraph
    :return: confirmation
    """
    global GLOBAL_DOC
    GLOBAL_DOC.add_paragraph(para)
    return "paragraph added!"

@tool
def write_header(header: str) -> str:
    """
    writes header in the doc
    :param header: header
    :return: confirmation
    """
    global GLOBAL_DOC
    p = GLOBAL_DOC.add_paragraph()
    run = p.add_run(header)
    run.font.size = Pt(14)
    run.bold = True
    return "header added!"

@tool
def write_sub_header(sub_header: str) -> str:
    """
    writes sub_header in the doc
    :param sub_header: sub header
    :return: confirmation
    """
    global GLOBAL_DOC
    p = GLOBAL_DOC.add_paragraph()
    run = p.add_run(sub_header)
    run.font.size = Pt(12)
    run.bold = True
    return "header added!"

@tool
def write_bullet_points(points: List[str]) -> str:
    """
    writes bullet points in the doc
    :param points: bullet points
    :return: confirmation
    """
    global GLOBAL_DOC
    for p in points:
        GLOBAL_DOC.add_paragraph(p, style = "List Bullet")
    return "bullet points added!"

@tool
def write_ordered_list(points: List[str]) -> str:
    """
    writes ordered list in the doc
    :param points: ordered list
    :return: confirmation
    """
    global GLOBAL_DOC
    for p in points:
        GLOBAL_DOC.add_paragraph(p, style="List Number")
    return "bullet points added!"

@tool
def save_doc(file_name: str) -> str:
    """
    Saves the document
    :param file_name: file name
    :return: confirmation
    """
    global GLOBAL_DOC

    if file_name and not file_name.endswith(".docx"):
        file_name = file_name+".docx"
    if not file_name:
        file_name = f"file_{str(np.random.randint(1,10_000))}.docx"

    GLOBAL_DOC.save(os.path.join("generated",file_name))
    GLOBAL_DOC = Document() # resetting the document to get fresh content
    return "Doc saved!"

@tool
def read_markdown() -> str:
    """
    Read instructions from a markdown file
    :return: instructions
    """
    with open(os.path.join("task","TODO.md"), "r") as md:
        content = md.read()
    return content

@tool
def invoke_rag(prompt: str) -> List[str]:
    """
    Invokes a RAG pipeline
    :param prompt: LLM generated prompt
    :return: retrieved context
    """
    results = vector_store.similarity_search(
        prompt,
        k=3
    )

    if not results:
        return []

    return [doc.page_content for doc in results]
